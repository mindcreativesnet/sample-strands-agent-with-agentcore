"""Document Manager for Markdown-based Document Creation

Manages markdown documents with outline tracking and DOCX conversion on finalization.
"""

import os
import base64
import io
from datetime import datetime
from typing import Optional, Dict, List, Any
from docx import Document
from docx.shared import RGBColor, Inches
import re


class DocumentManager:
    """Manages markdown document with outline and DOCX conversion"""

    def __init__(self, working_filepath: str = "/app/documents/current_document.md"):
        self.working_filepath = working_filepath
        self.docx_filepath = working_filepath.replace('.md', '.docx')
        self.markdown_content: str = ""
        self.metadata: Dict[str, Any] = {}
        self.outline: List[Dict[str, Any]] = []

        # Ensure documents directory exists
        os.makedirs(os.path.dirname(working_filepath), exist_ok=True)

    def create(self, title: str, outline: Optional[List[Dict[str, Any]]] = None) -> str:
        """Create new markdown document with title and optional outline"""

        # Store outline
        self.outline = outline if outline else []

        # Initialize markdown content with title
        self.markdown_content = f"# {title}\n\n"

        # Add outline structure as comments (for context, not rendered)
        if self.outline:
            self.markdown_content += "<!--\nDocument Outline:\n"
            self.markdown_content += self._format_outline_as_text(self.outline)
            self.markdown_content += "-->\n\n"

        self.metadata = {
            "title": title,
            "created_at": datetime.now().isoformat(),
            "has_outline": bool(self.outline)
        }

        self.auto_save()

        outline_info = ""
        if self.outline:
            outline_info = f"\n\nOutline:\n{self._format_outline_as_text(self.outline)}"

        return f"✅ Document '{title}' created and saved to {self.working_filepath}{outline_info}"

    def _format_outline_as_text(self, outline: List[Dict[str, Any]], indent: int = 0) -> str:
        """Format outline as readable text"""
        text = ""
        for item in outline:
            prefix = "  " * indent
            title = item.get('title', 'Untitled')
            text += f"{prefix}- {title}\n"

            if 'subsections' in item and item['subsections']:
                text += self._format_outline_as_text(item['subsections'], indent + 1)

        return text

    def get_outline(self) -> str:
        """Return current outline"""
        if not self.outline:
            return "No outline defined for this document."

        return "Document Outline:\n" + self._format_outline_as_text(self.outline)

    def write_content(self, content: str) -> str:
        """Append markdown content to document"""
        if not self.markdown_content:
            return "❌ No document created. Use create_document first."

        # Add content with proper spacing
        if not self.markdown_content.endswith('\n\n'):
            self.markdown_content += '\n\n'

        self.markdown_content += content
        self.auto_save()

        lines = content.count('\n') + 1
        chars = len(content)
        return f"✅ Added {lines} lines ({chars} characters) to document"

    def read_document(self) -> str:
        """Return current markdown content"""
        if not self.markdown_content:
            return "No document content available."

        return self.markdown_content

    def replace_content(self, find: str, replace: str) -> str:
        """Find and replace text in document"""
        if not self.markdown_content:
            return "❌ No document to edit."

        if find not in self.markdown_content:
            return f"❌ Text '{find}' not found in document."

        count = self.markdown_content.count(find)
        self.markdown_content = self.markdown_content.replace(find, replace)
        self.auto_save()

        return f"✅ Replaced {count} occurrence(s) of '{find}' with '{replace}'"

    def has_document(self) -> bool:
        """Check if document exists"""
        return bool(self.markdown_content)

    def clear(self) -> str:
        """Clear document and outline"""
        self.markdown_content = ""
        self.metadata = {}
        self.outline = []

        if os.path.exists(self.working_filepath):
            os.remove(self.working_filepath)
        if os.path.exists(self.docx_filepath):
            os.remove(self.docx_filepath)

        return "✅ Document cleared. Ready for new document."

    def auto_save(self) -> None:
        """Automatically save markdown to filesystem"""
        if self.markdown_content:
            with open(self.working_filepath, 'w', encoding='utf-8') as f:
                f.write(self.markdown_content)

    def get_info(self) -> str:
        """Return current document metadata"""
        if not self.markdown_content:
            return "No active document."

        lines = self.markdown_content.count('\n')
        chars = len(self.markdown_content)
        words = len(self.markdown_content.split())

        info = f"""Document Information:
- Title: {self.metadata.get('title', 'Untitled')}
- Created: {self.metadata.get('created_at', 'Unknown')}
- Lines: {lines}
- Words: {words}
- Characters: {chars}
- Has Outline: {'Yes' if self.outline else 'No'}
- Format: Markdown
- Saved to: {self.working_filepath}
"""

        if self.outline:
            info += f"\n{self.get_outline()}"

        return info

    def markdown_to_docx(self) -> Document:
        """Convert markdown content to DOCX document"""
        doc = Document()

        # Split content into lines for processing
        lines = self.markdown_content.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Skip HTML comments
            if line.startswith('<!--'):
                while i < len(lines) and '-->' not in lines[i]:
                    i += 1
                i += 1
                continue

            # Headings
            if line.startswith('# '):
                text = line[2:].strip()
                doc.add_heading(text, level=1)
            elif line.startswith('## '):
                text = line[3:].strip()
                doc.add_heading(text, level=2)
            elif line.startswith('### '):
                text = line[4:].strip()
                doc.add_heading(text, level=3)

            # Bullet lists
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                doc.add_paragraph(text, style='List Bullet')

            # Numbered lists
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                doc.add_paragraph(text, style='List Number')

            # Tables (markdown table syntax)
            elif line.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                i -= 1  # Back up one since we'll increment at end

                if len(table_lines) >= 2:  # Header + separator + data
                    self._add_table_to_doc(doc, table_lines)

            # Code blocks
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph(code_text)
                para.style = 'Normal'
                # Could add monospace font here

            # Blockquotes (for citations)
            elif line.startswith('> '):
                text = line[2:].strip()
                para = doc.add_paragraph(text)
                run = para.runs[0]
                run.italic = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

            # Images (markdown or base64 data URIs)
            elif line.startswith('!['):
                self._add_image_to_doc(doc, line)

            # Regular paragraphs
            elif line:
                # Handle inline formatting and citations
                self._add_formatted_paragraph(doc, line)

            i += 1

        return doc

    def _add_table_to_doc(self, doc: Document, table_lines: List[str]) -> None:
        """Add markdown table to DOCX document"""
        # Parse table
        rows = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
            rows.append(cells)

        # Skip separator line (usually second line with dashes)
        if len(rows) > 1 and all(cell.strip().replace('-', '').replace(':', '') == '' for cell in rows[1]):
            rows.pop(1)

        if not rows:
            return

        # Create table
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Fill table
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                if j < num_cols:  # Safety check
                    table.rows[i].cells[j].text = cell_data
                    # Make first row bold (header)
                    if i == 0:
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _add_formatted_paragraph(self, doc: Document, text: str) -> None:
        """Add paragraph with inline formatting (bold, italic, citations)"""
        para = doc.add_paragraph()

        # Handle citations in brackets [Author et al., Year]
        citation_pattern = r'\[([^\]]+(?:et al\.|[A-Z][a-z]+)[^\]]*(?:19|20)\d{2}[^\]]*)\]'

        parts = []
        last_end = 0

        for match in re.finditer(citation_pattern, text):
            # Add text before citation
            if match.start() > last_end:
                parts.append(('normal', text[last_end:match.start()]))

            # Add citation
            parts.append(('citation', match.group(0)))
            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            parts.append(('normal', text[last_end:]))

        # If no citations found, just add as normal text
        if not parts:
            parts = [('normal', text)]

        # Add runs with appropriate formatting
        for format_type, content in parts:
            # Handle bold **text** and italic *text* in normal text
            if format_type == 'normal':
                self._add_text_with_inline_formatting(para, content)
            else:  # citation
                run = para.add_run(content)
                run.italic = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    def _add_text_with_inline_formatting(self, para, text: str) -> None:
        """Add text with markdown inline formatting (bold, italic)"""
        # Parse **bold** and *italic* patterns
        # We process bold first, then italic within each segment
        import re

        # Pattern for **bold** - match ** followed by any chars (non-greedy) followed by **
        bold_pattern = r'\*\*(.+?)\*\*'

        # Process text segment by segment
        current_pos = 0

        for bold_match in re.finditer(bold_pattern, text):
            # Add text before bold
            if bold_match.start() > current_pos:
                before_text = text[current_pos:bold_match.start()]
                self._add_text_with_italic(para, before_text)

            # Add bold text
            bold_text = bold_match.group(1)
            run = para.add_run(bold_text)
            run.bold = True

            current_pos = bold_match.end()

        # Add remaining text
        if current_pos < len(text):
            remaining = text[current_pos:]
            self._add_text_with_italic(para, remaining)

    def _add_text_with_italic(self, para, text: str) -> None:
        """Add text with italic formatting (*text*)"""
        import re

        # Pattern for *italic* - must not be preceded or followed by another asterisk
        # This avoids matching the asterisks in **bold**
        italic_pattern = r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)'
        current_pos = 0

        for italic_match in re.finditer(italic_pattern, text):
            # Add text before italic
            if italic_match.start() > current_pos:
                para.add_run(text[current_pos:italic_match.start()])

            # Add italic text
            italic_text = italic_match.group(1)
            run = para.add_run(italic_text)
            run.italic = True

            current_pos = italic_match.end()

        # Add remaining text
        if current_pos < len(text):
            para.add_run(text[current_pos:])

    def _add_image_to_doc(self, doc: Document, line: str) -> None:
        """Add image to DOCX document from markdown image syntax

        Supports:
        - Base64 data URIs: ![title](data:image/png;base64,...)
        - Regular URLs/paths: ![title](path/to/image.png)
        """
        # Pattern: ![alt text](source)
        image_pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'
        match = re.search(image_pattern, line)

        if not match:
            return

        alt_text = match.group(1)
        image_source = match.group(2)

        try:
            # Check if it's a base64 data URI
            if image_source.startswith('data:image/'):
                # Extract base64 data
                # Format: data:image/png;base64,iVBORw0KGgoAAAANS...
                if ';base64,' in image_source:
                    base64_data = image_source.split(';base64,')[1]

                    # Decode base64 to bytes
                    image_bytes = base64.b64decode(base64_data)

                    # Create BytesIO stream
                    image_stream = io.BytesIO(image_bytes)

                    # Add image to document
                    # Set width to 6 inches (can be adjusted)
                    doc.add_picture(image_stream, width=Inches(6.0))

                    # Add caption if alt text exists
                    if alt_text:
                        caption = doc.add_paragraph(alt_text)
                        caption.alignment = 1  # Center alignment
                        # Make caption italic
                        for run in caption.runs:
                            run.italic = True

                    print(f"✅ Embedded base64 image: {alt_text or 'untitled'} ({len(image_bytes)} bytes)")
                else:
                    print(f"⚠️ Invalid base64 data URI format: {image_source[:50]}...")
            else:
                # Regular file path - attempt to load from filesystem
                if os.path.exists(image_source):
                    doc.add_picture(image_source, width=Inches(6.0))

                    if alt_text:
                        caption = doc.add_paragraph(alt_text)
                        caption.alignment = 1
                        for run in caption.runs:
                            run.italic = True

                    print(f"✅ Embedded image from file: {image_source}")
                else:
                    print(f"⚠️ Image file not found: {image_source}")
                    # Add placeholder text
                    para = doc.add_paragraph(f"[Image: {alt_text or image_source}]")
                    para.italic = True

        except Exception as e:
            print(f"❌ Error adding image: {e}")
            # Add error placeholder
            para = doc.add_paragraph(f"[Image Error: {alt_text or 'untitled'}]")
            para.italic = True
